from docx import Document as initDocument
from docx.document import Document
from docx.table import Table as WordTable
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches
from sqlalchemy.ext.asyncio import AsyncSession

from app.services.table_rep_manager import table_rep_manager as tableRepManager
from app.schemas.text_reports import TableStandart


def is_difficult_table_header(header: list[str]) -> bool:
    for h in header:
        if "|" in h:
            return True
    return False


def add_table_header(word_table: WordTable, header: list[str], idh: bool) -> None:
    if idh:
        cell_to_merge = None
        merging_cell_data = ""
        for i, h in enumerate(header):
            cell = word_table.cell(0, i)
            if "|" in h:
                row1, row2 = h.split(" | ")[:2]
                if row1 == merging_cell_data:
                    if cell_to_merge:
                        cell_to_merge.merge(word_table.cell(0, i))
                else:
                    cell_to_merge = cell
                    merging_cell_data = row1
                    cell.text = row1
                    cell.paragraphs[0].runs[0].bold = True
                word_table.cell(1, i).text = row2
            else:
                cell.text = h
                cell.paragraphs[0].runs[0].bold = True
                cell.merge(word_table.cell(1, i))
        return
    
    hdr_cells = word_table.rows[0].cells
    for i, header in enumerate(header):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True


def add_table_body(word_table: WordTable, body: list[list[str | int | float]], header_len: int) -> None:
    print("body data:", body[0])
    for row_idx, row_data in enumerate(body, header_len):
        row_cells = word_table.rows[row_idx].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)


def add_table_titels(doc: Document, section_name: str, table_name: str) -> None:
    section_text = doc.add_paragraph()
    section_run = section_text.add_run(section_name)
    section_run.bold = True
    
    title_text = doc.add_paragraph()
    title_run = title_text.add_run(table_name)
    title_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_run.italic = True


def create_table(doc: Document, rows: int, cols: int) -> WordTable:
    print("rows, cols:", rows, cols)
    word_table = doc.add_table(rows=rows, cols=cols)
    word_table.style = "Table Grid"
    return word_table


def add_table(doc: Document, table: TableStandart, section_name: str, table_name: str) -> None:
    add_table_titels(doc, section_name, table_name)
    
    is_dif_head = is_difficult_table_header(table.column_names)
    header_len = 2 if is_dif_head else 1
    
    word_table = create_table(doc, len(table.data) + header_len, len(table.column_names))
    
    add_table_header(word_table, table.column_names, is_dif_head)
    print(section_name, table_name)
    add_table_body(word_table, table.data, header_len)


def create_doc_file() -> Document:
    doc = initDocument()
    
    section = doc.sections[0]
    
    section.orientation = WD_ORIENTATION.LANDSCAPE
    
    section.page_width, section.page_height = section.page_height, section.page_width
    
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    return doc


async def get_docx_file(session: AsyncSession, section_num: int, exam_year: int, exam_type_id:int, subject_id: int):
    doc = create_doc_file()
    
    table_ind = 0
    while True:
        table_ind += 1
        table = await tableRepManager.get_table_by_section(session, section_num, table_ind, year=exam_year, exam_type_id=exam_type_id, subject_id=subject_id)
        if table.table_name == "Not found section or table":
            break
        
        add_table(doc, table, f"{section_num}.{table_ind} {table.table_name}", f"Таблица 2-{table_ind}")
    
    return doc